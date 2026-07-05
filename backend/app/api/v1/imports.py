"""
Data import routes:

  POST /import/preview          — upload file, get validation preview (no writes)
  POST /import/run              — commit a previewed import job
  GET  /import/jobs             — list import jobs
  GET  /import/jobs/{job_id}    — get job details + errors
  GET  /import/template/{fmt}   — download a blank template (csv/xlsx/docx)
"""
import uuid
from datetime import datetime, timezone
from typing import List, Optional

from fastapi import APIRouter, Depends, File, Form, HTTPException, Request, UploadFile
from fastapi.responses import Response
from pydantic import BaseModel
from sqlalchemy.orm import Session

from app.database import get_db
from app.models.app_user import AppUser
from app.models.import_job import ImportJob, ImportSource, ImportStatus
from app.core.dependencies import require_hr_or_above
from app.core.audit import log_action
from app.services.importer.parsers import detect_format, parse_csv, parse_xlsx, parse_docx
from app.services.importer.columns import normalise_row
from app.services.importer.validator import validate_rows
from app.services.importer.importer import run_import

router = APIRouter(prefix="/import", tags=["import"])

MAX_FILE_BYTES = 10 * 1024 * 1024   # 10 MB

_PARSERS = {"csv": parse_csv, "xlsx": parse_xlsx, "docx": parse_docx}
_SOURCE_ENUM = {"csv": ImportSource.csv, "xlsx": ImportSource.xlsx, "docx": ImportSource.docx}

MAGIC_BYTES = {
    "xlsx": [b"PK\x03\x04"],               # ZIP container (Office Open XML)
    "docx": [b"PK\x03\x04"],
    "csv":  [],                             # plain text — no magic bytes
}


from app.core.request_ip import client_ip as _client_ip


def _validate_magic(fmt: str, data: bytes) -> None:
    for magic in MAGIC_BYTES.get(fmt, []):
        if not data.startswith(magic):
            raise HTTPException(status_code=422, detail=f"File content does not match .{fmt} format.")


# ── Preview ───────────────────────────────────────────────────────────────────

@router.post("/preview")
async def preview_import(
    request: Request,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: AppUser = Depends(require_hr_or_above),
):
    """
    Upload a CSV/XLSX/DOCX file.
    Returns a validation summary without writing any person records.
    Saves an ImportJob in 'review' status so HR can decide whether to commit.
    """
    data = await file.read()
    if len(data) > MAX_FILE_BYTES:
        raise HTTPException(status_code=413, detail="File exceeds 10 MB limit.")
    if not data:
        raise HTTPException(status_code=422, detail="Uploaded file is empty.")

    fmt = detect_format(file.filename or "", file.content_type or "")
    _validate_magic(fmt, data)

    try:
        raw_rows = _PARSERS[fmt](data)
    except Exception as exc:
        raise HTTPException(status_code=422, detail=f"Could not parse file: {exc}")

    if not raw_rows:
        raise HTTPException(status_code=422, detail="No data rows found in the file.")

    normalised = [normalise_row(r) for r in raw_rows]
    results = validate_rows(normalised)

    valid_count = sum(1 for r in results if r.is_valid)
    invalid_count = len(results) - valid_count

    # Collect first 50 error rows for preview
    error_preview = [
        {"row": r.row_number, "errors": r.errors, "raw": r.raw}
        for r in results if not r.is_valid
    ][:50]

    # Sample of valid rows for display
    valid_preview = [
        r.cleaned for r in results if r.is_valid
    ][:10]

    job = ImportJob(
        source_type=_SOURCE_ENUM[fmt],
        filename=file.filename,
        status=ImportStatus.review,
        records_found=len(results),
        started_by=current_user.id,
        preview_data={
            "valid_count": valid_count,
            "invalid_count": invalid_count,
            "valid_sample": _serialise_preview(valid_preview),
            "errors": error_preview,
        },
    )
    db.add(job)
    db.commit()
    db.refresh(job)

    log_action(db, "import_previewed", user_id=str(current_user.id),
               detail={"filename": file.filename, "valid": valid_count, "invalid": invalid_count},
               ip_address=_client_ip(request))
    db.commit()

    return {
        "job_id": str(job.id),
        "filename": file.filename,
        "format": fmt,
        "records_found": len(results),
        "valid_count": valid_count,
        "invalid_count": invalid_count,
        "valid_sample": _serialise_preview(valid_preview),
        "errors": error_preview,
        "status": job.status.value,
    }


# ── Run / commit ──────────────────────────────────────────────────────────────

@router.post("/run")
async def run_import_job(
    request: Request,
    file: UploadFile = File(...),
    job_id: str = Form(...),
    main_company_id: str = Form(...),
    skip_errors: bool = Form(True),
    db: Session = Depends(get_db),
    current_user: AppUser = Depends(require_hr_or_above),
):
    """
    Commit an import: re-parses the file (client re-uploads) and writes records.
    Requires a job_id from /import/preview so we can record against the job.
    """
    try:
        job_uuid = uuid.UUID(job_id)
    except ValueError:
        raise HTTPException(status_code=422, detail="Invalid job_id.")

    job = db.query(ImportJob).filter(ImportJob.id == job_uuid).first()
    if not job:
        raise HTTPException(status_code=404, detail="Import job not found.")
    if job.status not in (ImportStatus.review, ImportStatus.failed):
        raise HTTPException(status_code=409, detail=f"Job is in '{job.status.value}' state and cannot be re-run.")

    data = await file.read()
    if len(data) > MAX_FILE_BYTES:
        raise HTTPException(status_code=413, detail="File exceeds 10 MB limit.")

    fmt = detect_format(file.filename or "", file.content_type or "")
    _validate_magic(fmt, data)

    try:
        raw_rows = _PARSERS[fmt](data)
    except Exception as exc:
        raise HTTPException(status_code=422, detail=f"Could not parse file: {exc}")

    normalised = [normalise_row(r) for r in raw_rows]
    results = validate_rows(normalised)

    if not skip_errors:
        invalid = [r for r in results if not r.is_valid]
        if invalid:
            raise HTTPException(
                status_code=422,
                detail=f"{len(invalid)} rows have validation errors. "
                       "Fix them or set skip_errors=true to import valid rows only.",
            )

    job = run_import(
        db=db,
        job=job,
        valid_rows=results,
        main_company_id=main_company_id,
        performed_by_id=str(current_user.id),
    )
    db.commit()

    log_action(db, "import_completed", user_id=str(current_user.id),
               detail={
                   "filename": file.filename,
                   "imported": job.records_imported,
                   "skipped": job.records_skipped,
               },
               ip_address=_client_ip(request))
    db.commit()

    return {
        "job_id": str(job.id),
        "status": job.status.value,
        "records_found": job.records_found,
        "records_imported": job.records_imported,
        "records_skipped": job.records_skipped,
        "errors": job.errors,
        "completed_at": job.completed_at.isoformat() if job.completed_at else None,
    }


# ── Job list / detail ─────────────────────────────────────────────────────────

@router.get("/jobs")
def list_import_jobs(
    db: Session = Depends(get_db),
    current_user: AppUser = Depends(require_hr_or_above),
    limit: int = 50,
    offset: int = 0,
):
    jobs = (
        db.query(ImportJob)
        .order_by(ImportJob.created_at.desc())
        .offset(offset)
        .limit(min(limit, 200))
        .all()
    )
    return [_job_summary(j) for j in jobs]


@router.get("/jobs/{job_id}")
def get_import_job(
    job_id: str,
    db: Session = Depends(get_db),
    current_user: AppUser = Depends(require_hr_or_above),
):
    try:
        jid = uuid.UUID(job_id)
    except ValueError:
        raise HTTPException(status_code=422, detail="Invalid job_id.")
    job = db.query(ImportJob).filter(ImportJob.id == jid).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found.")
    return {**_job_summary(job), "errors": job.errors, "preview_data": job.preview_data}


# ── Template download ─────────────────────────────────────────────────────────

@router.get("/template/{fmt}")
def download_template(
    fmt: str,
    current_user: AppUser = Depends(require_hr_or_above),
):
    fmt = fmt.lower()
    if fmt == "csv":
        content = _csv_template()
        return Response(content=content, media_type="text/csv",
                        headers={"Content-Disposition": "attachment; filename=import_template.csv"})
    if fmt == "xlsx":
        content = _xlsx_template()
        return Response(content=content,
                        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        headers={"Content-Disposition": "attachment; filename=import_template.xlsx"})
    if fmt == "docx":
        content = _docx_template()
        return Response(content=content,
                        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        headers={"Content-Disposition": "attachment; filename=import_template.docx"})
    raise HTTPException(status_code=422, detail="Format must be csv, xlsx, or docx.")


# ── Helpers ───────────────────────────────────────────────────────────────────

TEMPLATE_HEADERS = [
    "first_name", "last_name", "middle_name", "email", "phone",
    "person_type", "job_title", "department", "floor",
    "start_date", "company_name", "nfc_uid",
]

TEMPLATE_EXAMPLE = [
    "Jane", "Smith", "", "jane.smith@example.com", "+44 7700 900123",
    "employee", "Software Engineer", "Engineering", "2",
    "2026-01-15", "", "",
]

TEMPLATE_EXAMPLE_CTR = [
    "Bob", "Builder", "", "bob@buildright.com", "+44 7700 900456",
    "contractor", "Site Manager", "Construction", "1",
    "2026-06-01", "BuildRight Ltd", "",
]


def _csv_template() -> bytes:
    import io, csv
    buf = io.StringIO()
    w = csv.writer(buf)
    w.writerow(TEMPLATE_HEADERS)
    w.writerow(TEMPLATE_EXAMPLE)
    w.writerow(TEMPLATE_EXAMPLE_CTR)
    return buf.getvalue().encode("utf-8-sig")


def _xlsx_template() -> bytes:
    import io, openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "People Import"
    ws.append(TEMPLATE_HEADERS)
    ws.append(TEMPLATE_EXAMPLE)
    ws.append(TEMPLATE_EXAMPLE_CTR)
    # Bold headers
    from openpyxl.styles import Font
    for cell in ws[1]:
        cell.font = Font(bold=True)
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _docx_template() -> bytes:
    import io, docx
    from docx.shared import Pt
    document = docx.Document()
    document.add_heading("Employee Import Template", 1)
    document.add_paragraph(
        "Fill in the table below and import via Settings → Data Import. "
        "person_type must be 'employee' or 'contractor'. "
        "Contractors must have a company_name."
    )
    table = document.add_table(rows=3, cols=len(TEMPLATE_HEADERS))
    table.style = "Table Grid"
    for i, h in enumerate(TEMPLATE_HEADERS):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
    for i, v in enumerate(TEMPLATE_EXAMPLE):
        table.rows[1].cells[i].text = v
    for i, v in enumerate(TEMPLATE_EXAMPLE_CTR):
        table.rows[2].cells[i].text = v
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def _job_summary(job: ImportJob) -> dict:
    return {
        "id": str(job.id),
        "source_type": job.source_type.value,
        "filename": job.filename,
        "status": job.status.value,
        "records_found": job.records_found,
        "records_imported": job.records_imported,
        "records_skipped": job.records_skipped,
        "created_at": job.created_at.isoformat() if job.created_at else None,
        "completed_at": job.completed_at.isoformat() if job.completed_at else None,
    }


def _serialise_preview(rows: list) -> list:
    """Make date objects JSON-serialisable."""
    import datetime
    result = []
    for row in rows:
        clean = {}
        for k, v in row.items():
            clean[k] = v.isoformat() if isinstance(v, (datetime.date, datetime.datetime)) else v
        result.append(clean)
    return result
