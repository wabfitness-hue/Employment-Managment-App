"""
File parsers: CSV, XLSX, DOCX → list of raw dicts (one per row/table-row).
Each parser returns raw string values; normalisation happens later.
"""
import csv
import io
from typing import List


def parse_csv(file_bytes: bytes) -> List[dict]:
    """Parse CSV bytes → list of row dicts keyed by header row values."""
    text = file_bytes.decode("utf-8-sig")   # strips BOM if present
    reader = csv.DictReader(io.StringIO(text))
    rows = []
    for row in reader:
        # Skip entirely blank rows
        if any(v.strip() for v in row.values() if v):
            rows.append(dict(row))
    return rows


def parse_xlsx(file_bytes: bytes) -> List[dict]:
    """Parse first sheet of an XLSX workbook → list of row dicts."""
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    ws = wb.active
    rows_iter = ws.iter_rows(values_only=True)

    # First non-empty row = headers
    headers = None
    for row in rows_iter:
        if any(cell is not None for cell in row):
            headers = [str(h).strip() if h is not None else "" for h in row]
            break

    if not headers:
        return []

    result = []
    for row in rows_iter:
        values = [str(v).strip() if v is not None else "" for v in row]
        if any(values):
            result.append(dict(zip(headers, values)))

    wb.close()
    return result


def parse_docx(file_bytes: bytes) -> List[dict]:
    """
    Parse first table found in a DOCX file → list of row dicts.
    First row of the table is treated as headers.
    """
    import docx
    doc = docx.Document(io.BytesIO(file_bytes))

    if not doc.tables:
        raise ValueError("No tables found in the DOCX file. "
                         "Please format your data as a table (Insert → Table).")

    table = doc.tables[0]
    if len(table.rows) < 2:
        raise ValueError("Table in DOCX has fewer than 2 rows (need header + data).")

    headers = [cell.text.strip() for cell in table.rows[0].cells]
    result = []
    for row in table.rows[1:]:
        values = [cell.text.strip() for cell in row.cells]
        if any(values):
            result.append(dict(zip(headers, values)))

    return result


def detect_format(filename: str, content_type: str) -> str:
    """Determine file format from filename extension or MIME type."""
    name_lower = filename.lower()
    if name_lower.endswith(".csv"):
        return "csv"
    if name_lower.endswith((".xlsx", ".xls")):
        return "xlsx"
    if name_lower.endswith(".docx"):
        return "docx"
    # Fallback to content-type
    if "csv" in content_type:
        return "csv"
    if "spreadsheet" in content_type or "excel" in content_type:
        return "xlsx"
    if "wordprocessing" in content_type or "msword" in content_type:
        return "docx"
    raise ValueError(f"Unsupported file format: {filename!r} ({content_type})")
