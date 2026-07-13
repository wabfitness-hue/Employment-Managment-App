"""
Phase 7 Tests — Data import (CSV, XLSX, DOCX) + setup wizard.
"""
import csv
import io
import uuid
from datetime import date

import pytest

from app.services.importer.columns import normalise_header, normalise_row, normalise_person_type
from app.services.importer.parsers import parse_csv, parse_xlsx, parse_docx, detect_format
from app.services.importer.validator import validate_rows, RowResult


# ── Column normalisation ──────────────────────────────────────────────────────

class TestColumnNormalisation:
    def test_canonical_maps_to_itself(self):
        assert normalise_header("first_name") == "first_name"

    def test_alias_recognised(self):
        assert normalise_header("First Name") == "first_name"
        assert normalise_header("Surname") == "last_name"
        assert normalise_header("E-mail") == "email"
        assert normalise_header("Title") == "job_title"
        assert normalise_header("Dept") == "department"

    def test_unknown_header_returns_none(self):
        assert normalise_header("Unknown Column XYZ") is None

    def test_case_insensitive(self):
        assert normalise_header("FIRST NAME") == "first_name"
        assert normalise_header("first name") == "first_name"

    def test_normalise_row_drops_unknown_cols(self):
        row = {"First Name": "Jane", "Unknown": "ignored", "Email": "jane@x.com"}
        result = normalise_row(row)
        assert "first_name" in result
        assert "Unknown" not in result
        assert "unknown" not in result

    def test_normalise_row_strips_whitespace(self):
        row = {"first_name": "  Jane  "}
        result = normalise_row(row)
        assert result["first_name"] == "Jane"


class TestPersonTypeNormalisation:
    def test_employee_variants(self):
        for val in ["employee", "Employee", "EMPLOYEE", "staff", "permanent"]:
            assert normalise_person_type(val) == "employee"

    def test_contractor_variants(self):
        for val in ["contractor", "Contractor", "CTR", "temp", "freelance"]:
            assert normalise_person_type(val) == "contractor"

    def test_unknown_returns_none(self):
        assert normalise_person_type("intern") is None


# ── CSV parser ────────────────────────────────────────────────────────────────

def _make_csv(rows: list[list]) -> bytes:
    buf = io.StringIO()
    w = csv.writer(buf)
    for r in rows:
        w.writerow(r)
    return buf.getvalue().encode("utf-8")


HEADERS = ["first_name", "last_name", "email", "person_type",
           "job_title", "department", "start_date", "company_name"]


class TestCSVParser:
    def test_parses_basic_csv(self):
        data = _make_csv([
            HEADERS,
            ["Jane", "Smith", "jane@x.com", "employee", "Engineer", "Eng", "2026-01-01", ""],
        ])
        rows = parse_csv(data)
        assert len(rows) == 1
        assert rows[0]["first_name"] == "Jane"

    def test_skips_blank_rows(self):
        data = _make_csv([HEADERS, ["", "", "", "", "", "", "", ""], ["Bob", "Jones", "b@x.com", "employee", "Dev", "IT", "", ""]])
        rows = parse_csv(data)
        assert len(rows) == 1

    def test_handles_bom(self):
        with_bom = b"\xef\xbb\xbf" + _make_csv([HEADERS, ["Alice", "A", "a@x.com", "employee", "HR", "HR", "", ""]])
        rows = parse_csv(with_bom)
        assert rows[0]["first_name"] == "Alice"

    def test_multiple_rows(self):
        data = _make_csv([
            HEADERS,
            ["Alice", "A", "a@x.com", "employee", "HR", "HR", "", ""],
            ["Bob", "B", "b@x.com", "contractor", "Dev", "IT", "", "BuildCo"],
        ])
        assert len(parse_csv(data)) == 2

    def test_empty_csv_returns_empty(self):
        data = _make_csv([HEADERS])
        assert parse_csv(data) == []


# ── XLSX parser ───────────────────────────────────────────────────────────────

def _make_xlsx(rows: list[list]) -> bytes:
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    for r in rows:
        ws.append(r)
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


class TestXLSXParser:
    def test_parses_basic_xlsx(self):
        data = _make_xlsx([
            HEADERS,
            ["Jane", "Smith", "jane@x.com", "employee", "Engineer", "Eng", "2026-01-01", ""],
        ])
        rows = parse_xlsx(data)
        assert len(rows) == 1
        assert rows[0]["first_name"] == "Jane"

    def test_multiple_rows(self):
        data = _make_xlsx([
            HEADERS,
            ["Alice", "A", "a@x.com", "employee", "HR", "HR", "", ""],
            ["Bob", "B", "b@x.com", "contractor", "Dev", "IT", "", "BuildCo"],
        ])
        assert len(parse_xlsx(data)) == 2

    def test_skips_empty_rows(self):
        data = _make_xlsx([HEADERS, ["", "", "", "", "", "", "", ""], ["Alice", "A", "a@x.com", "employee", "HR", "HR", "", ""]])
        rows = parse_xlsx(data)
        assert len(rows) == 1


# ── DOCX parser ───────────────────────────────────────────────────────────────

def _make_docx(rows: list[list]) -> bytes:
    import docx
    doc = docx.Document()
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx].cells[c_idx].text = str(val)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestDOCXParser:
    def test_parses_table(self):
        data = _make_docx([
            HEADERS,
            ["Jane", "Smith", "jane@x.com", "employee", "Engineer", "Eng", "2026-01-01", ""],
        ])
        rows = parse_docx(data)
        assert len(rows) == 1
        assert rows[0]["first_name"] == "Jane"

    def test_no_table_raises(self):
        import docx as _docx
        doc = _docx.Document()
        doc.add_paragraph("No table here")
        buf = io.BytesIO()
        doc.save(buf)
        with pytest.raises(ValueError, match="No tables"):
            parse_docx(buf.getvalue())

    def test_multiple_rows(self):
        data = _make_docx([
            HEADERS,
            ["Alice", "A", "a@x.com", "employee", "HR", "HR", "", ""],
            ["Bob", "B", "b@x.com", "contractor", "Dev", "IT", "", "BuildCo"],
        ])
        assert len(parse_docx(data)) == 2


# ── Format detection ──────────────────────────────────────────────────────────

class TestFormatDetection:
    def test_csv_by_extension(self):
        assert detect_format("data.csv", "") == "csv"

    def test_xlsx_by_extension(self):
        assert detect_format("people.xlsx", "") == "xlsx"

    def test_docx_by_extension(self):
        assert detect_format("import.docx", "") == "docx"

    def test_csv_by_mimetype(self):
        assert detect_format("file", "text/csv") == "csv"

    def test_unsupported_raises(self):
        with pytest.raises(ValueError):
            detect_format("file.pdf", "application/pdf")


# ── Validator ─────────────────────────────────────────────────────────────────

def _valid_employee(**kwargs) -> dict:
    row = {
        "first_name": "Jane", "last_name": "Smith",
        "email": "jane@example.com", "person_type": "employee",
        "job_title": "Engineer", "department": "Engineering",
    }
    row.update(kwargs)
    return row


def _valid_contractor(**kwargs) -> dict:
    row = {
        "first_name": "Bob", "last_name": "Builder",
        "email": "bob@buildco.com", "person_type": "contractor",
        "job_title": "Site Manager", "department": "Construction",
        "company_name": "BuildCo Ltd",
    }
    row.update(kwargs)
    return row


class TestValidator:
    def test_valid_employee_passes(self):
        results = validate_rows([_valid_employee()])
        assert results[0].is_valid

    def test_valid_contractor_passes(self):
        results = validate_rows([_valid_contractor()])
        assert results[0].is_valid

    def test_missing_email_fails(self):
        row = _valid_employee()
        del row["email"]
        results = validate_rows([row])
        assert not results[0].is_valid
        assert any("email" in e for e in results[0].errors)

    def test_missing_first_name_fails(self):
        row = _valid_employee()
        del row["first_name"]
        results = validate_rows([row])
        assert not results[0].is_valid

    def test_invalid_email_fails(self):
        results = validate_rows([_valid_employee(email="not-an-email")])
        assert not results[0].is_valid

    def test_invalid_person_type_fails(self):
        results = validate_rows([_valid_employee(person_type="intern")])
        assert not results[0].is_valid

    def test_employee_variants_accepted(self):
        for val in ["Employee", "EMPLOYEE", "staff"]:
            results = validate_rows([_valid_employee(person_type=val)])
            assert results[0].is_valid, f"Expected valid for person_type={val!r}"
            assert results[0].cleaned["person_type"] == "employee"

    def test_contractor_without_company_fails(self):
        row = _valid_contractor()
        del row["company_name"]
        results = validate_rows([row])
        assert not results[0].is_valid
        assert any("company" in e.lower() for e in results[0].errors)

    def test_valid_date_accepted(self):
        results = validate_rows([_valid_employee(start_date="2026-01-15")])
        assert results[0].is_valid
        assert results[0].cleaned["start_date"] == date(2026, 1, 15)

    def test_uk_date_format_accepted(self):
        results = validate_rows([_valid_employee(start_date="15/01/2026")])
        assert results[0].is_valid

    def test_invalid_date_fails(self):
        results = validate_rows([_valid_employee(start_date="not-a-date")])
        assert not results[0].is_valid

    def test_missing_date_defaults_to_today(self):
        results = validate_rows([_valid_employee()])
        assert results[0].is_valid
        assert results[0].cleaned["start_date"] == date.today()

    def test_email_lowercased(self):
        results = validate_rows([_valid_employee(email="JANE@EXAMPLE.COM")])
        assert results[0].cleaned["email"] == "jane@example.com"

    def test_row_number_reported(self):
        results = validate_rows([_valid_employee(), _valid_employee(email="bad")])
        assert results[1].row_number == 3   # row 2 = index 1 → row_number = 2+1

    def test_multiple_valid_rows(self):
        rows = [_valid_employee(), _valid_contractor()]
        results = validate_rows(rows)
        assert all(r.is_valid for r in results)

    def test_optional_fields_preserved(self):
        results = validate_rows([_valid_employee(floor="3", phone="+44 7700 900123", nfc_uid="A1B2C3D4")])
        assert results[0].is_valid
        assert results[0].cleaned["floor"] == "3"
        assert results[0].cleaned["nfc_uid"] == "A1B2C3D4"


# ── Setup wizard (via HTTP client) ────────────────────────────────────────────

@pytest.fixture
def hr_client(client):
    """`client` with the HR-admin auth dependency satisfied (for protected routes)."""
    from app.core.dependencies import require_hr_or_above
    from app.models.app_user import AppUser, UserRole
    client.app.dependency_overrides[require_hr_or_above] = lambda: AppUser(
        email="hr@test.com", display_name="HR", password_hash="x", role=UserRole.hr_admin,
    )
    yield client
    client.app.dependency_overrides.pop(require_hr_or_above, None)


class TestSetupWizard:
    def test_status_requires_setup_on_fresh_db(self, client):
        r = client.get("/api/v1/setup/status")
        assert r.status_code == 200
        data = r.json()
        assert data["setup_required"] is True

    def test_company_step_creates_company(self, client):
        r = client.post("/api/v1/setup/company", json={
            "name": "Test Corp",
            "card_background_colour": "#1E40AF",
            "card_text_colour": "#FFFFFF",
        })
        assert r.status_code == 200
        data = r.json()
        assert data["name"] == "Test Corp"

    def test_company_blank_name_rejected(self, client):
        r = client.post("/api/v1/setup/company", json={"name": "   "})
        assert r.status_code == 422

    def test_company_invalid_hex_rejected(self, client):
        r = client.post("/api/v1/setup/company", json={
            "name": "Corp", "card_background_colour": "blue"
        })
        assert r.status_code == 422

    def test_full_setup_flow(self, client):
        # Step 1: company
        r = client.post("/api/v1/setup/company", json={"name": "Acme Corp"})
        assert r.status_code == 200

        # Step 2: admin
        r = client.post("/api/v1/setup/admin", json={
            "email": "admin@acme.com",
            "password": "Str0ng!Pass#2026",
            "full_name": "System Administrator",
        })
        assert r.status_code == 200
        data = r.json()
        assert "mfa_provisioning_uri" in data
        assert "mfa_secret" in data
        admin_id = data["admin_id"]
        mfa_secret = data["mfa_secret"]

        # Step 3: complete with valid TOTP
        import pyotp
        token = pyotp.TOTP(mfa_secret).now()
        r = client.post("/api/v1/setup/complete", json={
            "admin_id": admin_id,
            "mfa_token": token,
        })
        assert r.status_code == 200
        data = r.json()
        assert "access_token" in data
        assert data["role"] == "super_admin"

        # Re-completing after setup must be rejected (no passwordless re-login).
        token2 = pyotp.TOTP(mfa_secret).now()
        r = client.post("/api/v1/setup/complete", json={
            "admin_id": admin_id,
            "mfa_token": token2,
        })
        assert r.status_code == 409

    def test_setup_blocked_after_completion(self, client):
        """After a super_admin exists, /setup/company must return 409."""
        import pyotp

        # Ensure we have a company then admin (may already exist from prior tests in same client)
        client.post("/api/v1/setup/company", json={"name": "BlockTest Corp"})
        r_admin = client.post("/api/v1/setup/admin", json={
            "email": "blocker@corp.com",
            "password": "Str0ng!Pass#2026",
            "full_name": "Block Admin",
        })
        if r_admin.status_code == 200:
            data = r_admin.json()
            token = pyotp.TOTP(data["mfa_secret"]).now()
            client.post("/api/v1/setup/complete", json={
                "admin_id": data["admin_id"], "mfa_token": token
            })

        # Whether we just created one or it already existed, setup should be blocked now
        r = client.post("/api/v1/setup/company", json={"name": "X"})
        assert r.status_code == 409

    def test_get_prefixes(self, hr_client):
        client = hr_client
        # Seed some prefixes first
        client.post("/api/v1/setup/prefixes", json={"prefixes": [
            {"prefix": "DIR", "label": "Director", "person_type": "employee"},
            {"prefix": "CTR", "label": "Contractor", "person_type": "contractor"},
        ]})
        r = client.get("/api/v1/setup/prefixes")
        assert r.status_code == 200
        prefixes = r.json()
        codes = [p["prefix"] for p in prefixes]
        assert "DIR" in codes
        assert "CTR" in codes

    def test_setup_prefixes(self, hr_client):
        client = hr_client
        r = client.post("/api/v1/setup/prefixes", json={"prefixes": [
            {"prefix": "DIR", "label": "Director", "person_type": "employee"},
            {"prefix": "CTR", "label": "Contractor", "person_type": "contractor"},
        ]})
        assert r.status_code == 200
        assert r.json()["created"] == 2

    def test_prefixes_require_auth(self, client):
        # Unauthenticated access to the destructive prefix routes must be rejected.
        r = client.post("/api/v1/setup/prefixes", json={"prefixes": [
            {"prefix": "DIR", "label": "Director", "person_type": "employee"},
        ]})
        assert r.status_code in (401, 403)
        assert client.get("/api/v1/setup/prefixes").status_code in (401, 403)

    def test_duplicate_prefix_rejected(self, hr_client):
        client = hr_client
        r = client.post("/api/v1/setup/prefixes", json={"prefixes": [
            {"prefix": "DIR", "label": "Director", "person_type": "employee"},
            {"prefix": "DIR", "label": "Director2", "person_type": "employee"},
        ]})
        assert r.status_code == 422

    def test_invalid_prefix_format_rejected(self, hr_client):
        client = hr_client
        r = client.post("/api/v1/setup/prefixes", json={"prefixes": [
            {"prefix": "X1", "label": "Bad", "person_type": "employee"},
        ]})
        assert r.status_code == 422

    def test_weak_password_rejected_in_admin_step(self, client):
        client.post("/api/v1/setup/company", json={"name": "Test Co"})
        r = client.post("/api/v1/setup/admin", json={
            "email": "admin@test.com",
            "password": "weak",
            "full_name": "Admin User",
        })
        assert r.status_code == 422


# ── Audit log viewer ──────────────────────────────────────────────────────────

@pytest.fixture
def it_client(client):
    """`client` with the IT-admin auth dependency satisfied (for the audit routes)."""
    from app.core.dependencies import require_it_or_above
    from app.models.app_user import AppUser, UserRole
    client.app.dependency_overrides[require_it_or_above] = lambda: AppUser(
        email="it@test.com", display_name="IT", password_hash="x", role=UserRole.it_admin,
    )
    yield client
    client.app.dependency_overrides.pop(require_it_or_above, None)


class TestAuditLog:
    def test_requires_auth(self, client):
        assert client.get("/api/v1/audit").status_code in (401, 403)

    def test_lists_entries(self, it_client):
        client = it_client
        # A failed login writes a "login_failed" audit entry.
        client.post("/api/v1/auth/login", json={"email": "nobody@example.com", "password": "wrong"})
        r = client.get("/api/v1/audit")
        assert r.status_code == 200
        data = r.json()
        assert data["total"] >= 1
        assert any(i["action"] == "login_failed" for i in data["items"])
        # newest-first
        ts = [i["timestamp"] for i in data["items"] if i["timestamp"]]
        assert ts == sorted(ts, reverse=True)

    def test_filter_by_action(self, it_client):
        client = it_client
        client.post("/api/v1/auth/login", json={"email": "nobody@example.com", "password": "wrong"})
        r = client.get("/api/v1/audit", params={"action": "login_failed"})
        assert r.status_code == 200
        assert all(i["action"] == "login_failed" for i in r.json()["items"])

    def test_actions_list(self, it_client):
        client = it_client
        client.post("/api/v1/auth/login", json={"email": "nobody@example.com", "password": "wrong"})
        r = client.get("/api/v1/audit/actions")
        assert r.status_code == 200
        assert "login_failed" in r.json()["actions"]


# ── People CSV export ─────────────────────────────────────────────────────────

@pytest.fixture
def any_client(client):
    from app.database import get_db
    from app.core.dependencies import require_any_role
    from app.models.app_user import AppUser, UserRole
    # Insert the acting user so the export's audit-log FK (user_id) is satisfied.
    db = client.app.dependency_overrides[get_db]()
    user = AppUser(id=uuid.uuid4(), email="hr-exp@test.com", display_name="HR",
                   password_hash="x", role=UserRole.hr_admin)
    db.add(user); db.commit()
    client.app.dependency_overrides[require_any_role] = lambda: user
    yield client
    client.app.dependency_overrides.pop(require_any_role, None)


class TestPeopleExport:
    def test_requires_auth(self, client):
        assert client.get("/api/v1/people/export.csv").status_code in (401, 403)

    def test_returns_csv_with_header(self, any_client):
        r = any_client.get("/api/v1/people/export.csv")
        assert r.status_code == 200
        assert "text/csv" in r.headers["content-type"]
        assert "attachment" in r.headers.get("content-disposition", "")
        assert r.text.splitlines()[0].startswith("Employee ID,First Name,Last Name,Email")

    def test_csv_safe_neutralises_formula_triggers(self):
        from app.api.v1.people import _csv_safe
        assert _csv_safe("=cmd|'/c calc'!A1") == "'=cmd|'/c calc'!A1"
        assert _csv_safe("+1+1") == "'+1+1"
        assert _csv_safe("-1") == "'-1"
        assert _csv_safe("@SUM(1,1)") == "'@SUM(1,1)"
        assert _csv_safe("Engineering") == "Engineering"  # untouched
        assert _csv_safe("") == ""


# ── Printers CRUD ─────────────────────────────────────────────────────────────

@pytest.fixture
def hr_client2(client):
    """A second HR-scoped client fixture (own name to avoid collisions), for
    the printers tests — user is inserted so audit-log FK writes succeed."""
    from app.database import get_db
    from app.core.dependencies import require_hr_or_above, require_any_role
    from app.models.app_user import AppUser, UserRole
    db = client.app.dependency_overrides[get_db]()
    user = AppUser(id=uuid.uuid4(), email="hr-printers@test.com", display_name="HR",
                   password_hash="x", role=UserRole.hr_admin)
    db.add(user); db.commit()
    client.app.dependency_overrides[require_hr_or_above] = lambda: user
    client.app.dependency_overrides[require_any_role] = lambda: user
    yield client
    client.app.dependency_overrides.pop(require_hr_or_above, None)
    client.app.dependency_overrides.pop(require_any_role, None)


class TestPrinters:
    def test_requires_auth_to_list(self, client):
        assert client.get("/api/v1/printers").status_code in (401, 403)

    def test_create_and_list_os_printer(self, hr_client2):
        client = hr_client2
        r = client.post("/api/v1/printers", json={
            "label": "3rd Floor Printer", "target_type": "os", "target": "HP LaserJet",
        })
        assert r.status_code == 200
        data = r.json()
        assert data["label"] == "3rd Floor Printer"
        assert data["target_type"] == "os"

        r2 = client.get("/api/v1/printers")
        assert r2.status_code == 200
        assert any(p["label"] == "3rd Floor Printer" for p in r2.json())

    def test_zebra_requires_valid_ip(self, hr_client2):
        client = hr_client2
        r = client.post("/api/v1/printers", json={
            "label": "Zebra Card Printer", "target_type": "zebra", "target": "not-an-ip",
        })
        assert r.status_code == 422

    def test_zebra_valid_ip_accepted(self, hr_client2):
        client = hr_client2
        r = client.post("/api/v1/printers", json={
            "label": "Zebra Card Printer", "target_type": "zebra", "target": "192.168.1.50",
        })
        assert r.status_code == 200
        assert r.json()["target"] == "192.168.1.50"

    def test_delete_printer(self, hr_client2):
        client = hr_client2
        r = client.post("/api/v1/printers", json={
            "label": "Temp Printer", "target_type": "os", "target": "Temp",
        })
        pid = r.json()["id"]
        r2 = client.delete(f"/api/v1/printers/{pid}")
        assert r2.status_code == 200
        r3 = client.get("/api/v1/printers")
        assert not any(p["id"] == pid for p in r3.json())

    def test_blank_label_rejected(self, hr_client2):
        client = hr_client2
        r = client.post("/api/v1/printers", json={
            "label": "   ", "target_type": "os", "target": "Something",
        })
        assert r.status_code == 422
